from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from generate_manual_usuario import BLUE, DARK, TEAL, add_table, configure_document


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Hoja_de_Firmas_Manual_PROMATEC_LUGARTH.docx"
LOGO = ROOT / "public" / "img" / "lugarth.png"


def build() -> Path:
    doc = Document()
    configure_document(doc)

    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(5)
    if LOGO.exists():
        header.add_run().add_picture(str(LOGO), width=Inches(0.72))

    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand.paragraph_format.space_after = Pt(2)
    run = brand.add_run("PROMATEC LUGARTH")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(DARK)

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version.paragraph_format.space_after = Pt(12)
    vr = version.add_run("Manual de Usuario | Versión 1.1 | Agosto 2026")
    vr.font.size = Pt(8)
    vr.font.color.rgb = RGBColor(82, 96, 109)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    tr = title.add_run("CONTROL DE ENTREGA Y APROBACIÓN")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.color.rgb = RGBColor.from_string(BLUE)

    intro = doc.add_paragraph(
        "Complete esta hoja después de colocar las capturas, revisar el contenido y validar la operación del sistema."
    )
    intro.paragraph_format.space_after = Pt(10)

    heading = doc.add_paragraph()
    heading.paragraph_format.space_after = Pt(5)
    hr = heading.add_run("Lista de entrega del manual")
    hr.bold = True
    hr.font.size = Pt(12)
    hr.font.color.rgb = RGBColor.from_string(DARK)

    checklist = [
        "Se sustituyeron los 43 marcadores web y los 4 marcadores móviles.",
        "No aparecen contraseñas ni datos personales o bancarios sin protección.",
        "Las capturas corresponden a la versión instalada en la computadora de operación.",
        "Se revisaron nombres de botones, periodos y roles.",
        "Se validó una impresión de recibos normales, diferencias IMSS y horas de alumnos.",
        "Se confirmó la operación con un usuario distinto de Administrador.",
        "El archivo final se guardó en Word y PDF.",
        "Se asignó responsable y fecha de próxima revisión.",
    ]
    for item in checklist:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.12)
        paragraph.paragraph_format.space_after = Pt(4)
        box = paragraph.add_run("[  ]  ")
        box.bold = True
        box.font.color.rgb = RGBColor.from_string(TEAL)
        paragraph.add_run(item)

    approval = doc.add_paragraph()
    approval.paragraph_format.space_before = Pt(10)
    approval.paragraph_format.space_after = Pt(6)
    ar = approval.add_run("Aprobación")
    ar.bold = True
    ar.font.size = Pt(12)
    ar.font.color.rgb = RGBColor.from_string(DARK)

    add_table(
        doc,
        ["Responsable", "Nombre", "Firma", "Fecha"],
        [
            ["Elaboró", "", "", ""],
            ["Revisó", "", "", ""],
            ["Autorizó", "", "", ""],
        ],
        [1800, 3000, 2760, 1800],
        font_size=9,
    )

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(18)
    fr = footer.add_run("DOCUMENTO DE USO INTERNO")
    fr.bold = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor.from_string(TEAL)

    doc.core_properties.title = "Hoja de Firmas - Manual PROMATEC LUGARTH"
    doc.core_properties.subject = "Control de entrega y aprobación del manual"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
