from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs"
OUTPUT_PATH = OUTPUT_DIR / "Manual_de_Usuario_Sistema_Nominas_PROMATEC_LUGARTH.docx"
LUGARTH_LOGO = ROOT / "public" / "img" / "lugarth.png"
PROMATEC_LOGO = ROOT / "public" / "img" / "promatec.png"

BLUE = "2E74B5"
DARK_BLUE = "17365D"
TEAL = "008A83"
LIGHT_BLUE = "EAF2F8"
LIGHT_TEAL = "E8F7F5"
LIGHT_AMBER = "FFF6DD"
LIGHT_RED = "FDECEC"
LIGHT_GRAY = "F5F7FA"
MID_GRAY = "D9E2EC"
DARK = "172B4D"
GRAY = "52606D"
WHITE = "FFFFFF"


CAPTURES = [
    {
        "id": "CAP-01",
        "screen": "Inicio de sesión",
        "state": "Pantalla completa, sin escribir una contraseña real.",
        "focus": "Logo, correo, contraseña, recordar sesión, recuperar contraseña y acceso a registro.",
    },
    {
        "id": "CAP-02",
        "screen": "Registro de nueva cuenta",
        "state": "Formulario vacío o con datos ficticios.",
        "focus": "Nombre, correo, contraseña, confirmación y botón de registro.",
    },
    {
        "id": "CAP-03",
        "screen": "Cuenta pendiente",
        "state": "Ingresar con una cuenta recién registrada todavía no aprobada.",
        "focus": "Mensaje que indica que el acceso requiere aprobación.",
    },
    {
        "id": "CAP-04",
        "screen": "Panel principal",
        "state": "Administrador, semana con información representativa.",
        "focus": "Indicadores, periodo contable, nómina, asistencia, personal y estado del sistema.",
    },
    {
        "id": "CAP-05",
        "screen": "Navegación general",
        "state": "Menú lateral abierto y barra superior visible.",
        "focus": "Módulos, usuario, rol, estado en línea, botón del menú y Salir.",
    },
    {
        "id": "CAP-06",
        "screen": "Empleados - alta",
        "state": "Formulario en modo Registrar con un ejemplo ficticio.",
        "focus": "Número, nombre, puesto, fecha de ingreso, forma de pago, sueldo/tarifa y préstamo.",
    },
    {
        "id": "CAP-07",
        "screen": "Empleados - más datos",
        "state": "Sección Más datos desplegada; usar valores ficticios.",
        "focus": "IMSS, ISR, INFONAVIT, banco, cuenta/CLABE y datos personales/contacto.",
    },
    {
        "id": "CAP-08",
        "screen": "Directorio de empleados",
        "state": "Filtro Activos; mostrar búsqueda y al menos dos filas.",
        "focus": "Foto, número, nombre, puesto, pago y acciones disponibles.",
    },
    {
        "id": "CAP-09",
        "screen": "Visor de fotografía",
        "state": "Abrir la foto de un empleado de prueba.",
        "focus": "Imagen ampliada, nombre/número y control para cerrar.",
    },
    {
        "id": "CAP-10",
        "screen": "Perfil del empleado",
        "state": "Empleado activo con datos de ejemplo.",
        "focus": "Encabezado, estado y pestañas Perfil, Nómina, Asistencia y Vacaciones.",
    },
    {
        "id": "CAP-11",
        "screen": "Perfil - nómina y préstamo",
        "state": "Pestaña Nómina de un empleado con valores representativos.",
        "focus": "Sueldo/tarifa, deducciones, forma de pago, deuda y descuento del préstamo.",
    },
    {
        "id": "CAP-12",
        "screen": "Perfil - asistencia y vacaciones",
        "state": "Pestaña con historial suficiente para mostrar los indicadores.",
        "focus": "Vacaciones totales, tomadas, restantes, faltas y fechas.",
    },
    {
        "id": "CAP-13",
        "screen": "Baja de empleado",
        "state": "Diálogo o formulario de baja abierto, sin confirmar.",
        "focus": "Fecha de baja, motivo, advertencia y botón de confirmación.",
    },
    {
        "id": "CAP-14",
        "screen": "Perfil de empleado dado de baja",
        "state": "Registro de prueba ya dado de baja.",
        "focus": "Fecha, motivo, días laborados totales y días laborados durante el año de la baja.",
    },
    {
        "id": "CAP-15",
        "screen": "Bajas y restauración",
        "state": "Filtro Bajas/Papelera con un empleado de prueba.",
        "focus": "Acción Restaurar y aviso de número ocupado, si aplica.",
    },
    {
        "id": "CAP-16",
        "screen": "Acceso móvil Mi Lugarth",
        "state": "Perfil de empleado; sección de acceso móvil desplegada.",
        "focus": "Estado, usuario, contraseña temporal, Guardar acceso y Desactivar.",
    },
    {
        "id": "CAP-17",
        "screen": "Asistencias - Captura y Reloj",
        "state": "Pestaña Captura, sin archivo seleccionado.",
        "focus": "Archivo CSV, Inicio semana, Fin semana y botón de previsualización/importación.",
    },
    {
        "id": "CAP-18",
        "screen": "Asistencias - captura manual",
        "state": "Empleado de prueba seleccionado.",
        "focus": "Empleado, fecha, tipo, entrada, salida, indicadores y Guardar.",
    },
    {
        "id": "CAP-19",
        "screen": "Asistencias - control semanal",
        "state": "Semana con registros normales, falta, incidencia e incompleta.",
        "focus": "Jueves a miércoles, búsqueda, cambio de semana, totales y acciones editar/eliminar.",
    },
    {
        "id": "CAP-20",
        "screen": "Revisión CSV - resumen",
        "state": "CSV previsualizado, antes de aprobar.",
        "focus": "Rango efectivo, total, seleccionadas, sin registro, incompletas, actualizan y sin empleado.",
    },
    {
        "id": "CAP-21",
        "screen": "Revisión CSV - fila editable",
        "state": "Filtro Incompletas y una celda abierta para corrección.",
        "focus": "Empleado, fecha, tipo, entrada, salida, selección y estado Incompleta.",
    },
    {
        "id": "CAP-22",
        "screen": "Revisión CSV - faltas",
        "state": "Filtro Faltas con una falta editable.",
        "focus": "Estado Sin registro, tipo Falta, fecha, selección y texto explicativo.",
    },
    {
        "id": "CAP-23",
        "screen": "Revisión CSV - aprobación",
        "state": "Parte inferior o encabezado donde aparezcan las acciones finales.",
        "focus": "Seleccionar todo, Quitar selección, Aprobar seleccionadas y Descartar revisión.",
    },
    {
        "id": "CAP-24",
        "screen": "Control de vacaciones",
        "state": "Pestaña Control Vacaciones con varios empleados.",
        "focus": "Totales, usadas, restantes, búsqueda y datos por empleado.",
    },
    {
        "id": "CAP-25",
        "screen": "Control de faltas",
        "state": "Pestaña Control Faltas, año seleccionado.",
        "focus": "Selector de año, total por empleado y desglose de fechas.",
    },
    {
        "id": "CAP-26",
        "screen": "Horas de alumnos",
        "state": "Semana con alumnos y horas registradas.",
        "focus": "Semana, búsqueda, selección, número de empleado, horas y botón PDF.",
    },
    {
        "id": "CAP-27",
        "screen": "PDF de horas de alumnos",
        "state": "Vista previa del PDF.",
        "focus": "Dos registros por hoja y número de empleado claramente visible.",
    },
    {
        "id": "CAP-28",
        "screen": "Nóminas - periodo y herramientas",
        "state": "Semana con nóminas existentes.",
        "focus": "Semana, estado, orden, banco, búsqueda, Excel global/IMSS y botones PDF.",
    },
    {
        "id": "CAP-29",
        "screen": "Nóminas - grupo bancario",
        "state": "Grupo de un banco desplegado.",
        "focus": "Nombre del banco, cuentas, Seleccionar grupo y PDF grupo.",
    },
    {
        "id": "CAP-30",
        "screen": "Nóminas - tarjeta del empleado",
        "state": "Empleado pendiente con asistencia completa.",
        "focus": "Estado, cuenta, pago neto, deuda, horas pendientes, Excel y Generar/Regenerar.",
    },
    {
        "id": "CAP-31",
        "screen": "Nóminas - préstamo y otros ajustes",
        "state": "Secciones visibles con valores de ejemplo controlados.",
        "focus": "Entregar hoy, Descontar hoy, descuento manual y días de vacaciones.",
    },
    {
        "id": "CAP-32",
        "screen": "Nóminas - Diferencia IMSS",
        "state": "Empleado con diferencia distinta de cero.",
        "focus": "Depósito IMSS, diferencia semanal, suma total y Guardar IMSS.",
    },
    {
        "id": "CAP-33",
        "screen": "Nóminas - ajustes avanzados",
        "state": "Panel de ajustes ampliado.",
        "focus": "Faltas, vacaciones, incapacidad, horas extra detectadas/por pagar y Guardar ajustes.",
    },
    {
        "id": "CAP-34",
        "screen": "Recibo individual de nómina",
        "state": "Vista previa PDF con datos ficticios o autorizados.",
        "focus": "Periodo, percepciones, deducciones, neto y firma.",
    },
    {
        "id": "CAP-35",
        "screen": "PDF masivo de nómina",
        "state": "Vista previa de una hoja completa.",
        "focus": "Dos recibos por hoja y separación correcta.",
    },
    {
        "id": "CAP-36",
        "screen": "PDF Diferencias IMSS",
        "state": "Vista previa de empleado con diferencia.",
        "focus": "Recibo independiente, concepto Sueldo, diferencia y demás importes en cero.",
    },
    {
        "id": "CAP-37",
        "screen": "Historial de recibos",
        "state": "Historial con más de una semana.",
        "focus": "Búsqueda, periodo, estado, empleado y acción PDF.",
    },
    {
        "id": "CAP-38",
        "screen": "Días festivos",
        "state": "Año actual con días oficiales y uno empresarial de prueba.",
        "focus": "Selector de año, Generar oficiales, alta manual, activo/inactivo y acciones.",
    },
    {
        "id": "CAP-39",
        "screen": "Base de datos",
        "state": "Pantalla principal, sin iniciar una restauración.",
        "focus": "Conexión, nombre, tablas, Exportar respaldo e Importar/Restaurar.",
    },
    {
        "id": "CAP-40",
        "screen": "Usuarios y permisos",
        "state": "Administrador; una cuenta pendiente y una activa.",
        "focus": "Rol, Aprobado, Deshabilitado, permisos, último acceso y Guardar.",
    },
    {
        "id": "CAP-41",
        "screen": "Auditoría",
        "state": "Registros variados; abrir un detalle no sensible.",
        "focus": "Filtros, usuario, acción, fecha, descripción, cambios y detalle técnico.",
    },
    {
        "id": "CAP-42",
        "screen": "Configuración del perfil",
        "state": "Cuenta de prueba o datos protegidos.",
        "focus": "Datos de perfil, cambio de contraseña y eliminación de cuenta cuando esté permitida.",
    },
    {
        "id": "APP-01",
        "screen": "Mi Lugarth - inicio de sesión",
        "state": "Aplicación móvil instalada; usar credenciales de prueba.",
        "focus": "Usuario, contraseña y acceso.",
    },
    {
        "id": "APP-02",
        "screen": "Mi Lugarth - inicio/resumen",
        "state": "Empleado de prueba vinculado.",
        "focus": "Identidad del empleado y resumen disponible.",
    },
    {
        "id": "APP-03",
        "screen": "Mi Lugarth - asistencias",
        "state": "Empleado con registros sincronizados.",
        "focus": "Fechas, entrada, salida, faltas/incidencias y horas extra de 0.5 cuando existan.",
    },
    {
        "id": "APP-04",
        "screen": "Mi Lugarth - nóminas pagadas",
        "state": "Empleado con al menos una nómina marcada como pagada.",
        "focus": "Semana, estado pagado, resumen e historial.",
    },
]


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_field(run, instruction: str) -> None:
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_begin, instr_text, fld_char_separate, fld_char_end])


def set_paragraph_box(paragraph, fill: str, border: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge in ("top", "left", "bottom", "right"):
        tag = p_bdr.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            p_bdr.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "8" if edge == "left" else "4")
        tag.set(qn("w:space"), "5")
        tag.set(qn("w:color"), border)
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), "100")
    spacing.set(qn("w:after"), "100")
    keep_lines = p_pr.find(qn("w:keepLines"))
    if keep_lines is None:
        keep_lines = OxmlElement("w:keepLines")
        p_pr.append(keep_lines)
    keep_lines.set(qn("w:val"), "true")


def keep_with_next(paragraph, value=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    element = p_pr.find(qn("w:keepNext"))
    if element is None:
        element = OxmlElement("w:keepNext")
        p_pr.append(element)
    element.set(qn("w:val"), "1" if value else "0")


def add_text(doc: Document, text: str, *, bold=False, color=None, style=None, align=None):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def new_numbering_id(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    style = doc.styles["List Number"]._element
    style_num_id = style.find("./w:pPr/w:numPr/w:numId", style.nsmap)
    base_num_id = int(style_num_id.get(qn("w:val"))) if style_num_id is not None else 5
    base_num = next(
        (num for num in numbering.findall(qn("w:num")) if int(num.get(qn("w:numId"))) == base_num_id),
        None,
    )
    if base_num is None:
        raise RuntimeError("No se encontró la numeración base de List Number.")

    abstract_num_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing_ids = [int(num.get(qn("w:numId"))) for num in numbering.findall(qn("w:num"))]
    next_num_id = max(existing_ids, default=0) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return next_num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_element = num_pr.find(qn("w:numId"))
    if num_id_element is None:
        num_id_element = OxmlElement("w:numId")
        num_pr.append(num_id_element)
    num_id_element.set(qn("w:val"), str(num_id))


def add_steps(doc: Document, title: str, items: list[str]) -> None:
    doc.add_heading(title, level=3)
    num_id = new_numbering_id(doc)
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        apply_numbering(paragraph, num_id)
        paragraph.add_run(item)


def add_note(doc: Document, title: str, text: str, kind="info") -> None:
    colors = {
        "info": (LIGHT_BLUE, BLUE),
        "good": (LIGHT_TEAL, TEAL),
        "warn": (LIGHT_AMBER, "C58A00"),
        "danger": (LIGHT_RED, "C62828"),
    }
    fill, border = colors[kind]
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"{title.upper()}: ")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(border)
    paragraph.add_run(text)
    set_paragraph_box(paragraph, fill, border)


def add_capture(doc: Document, capture_id: str) -> None:
    capture = next(item for item in CAPTURES if item["id"] == capture_id)
    heading = doc.add_paragraph()
    heading.style = doc.styles["Capture Label"]
    heading.add_run(f"{capture['id']} | {capture['screen']}")
    keep_with_next(heading)
    box = doc.add_paragraph()
    box.paragraph_format.left_indent = Inches(0.12)
    box.paragraph_format.right_indent = Inches(0.12)
    box.paragraph_format.space_before = Pt(0)
    box.paragraph_format.space_after = Pt(10)
    r1 = box.add_run("INSERTAR CAPTURA AQUÍ\n")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor.from_string(BLUE)
    r2 = box.add_run(f"Preparación: {capture['state']}\n")
    r2.bold = True
    box.add_run(f"Debe verse: {capture['focus']}\n")
    box.add_run("Recorte: conservar el título de la pantalla y los controles mencionados; ocultar datos sensibles.")
    set_paragraph_box(box, "F8FAFC", "9FB3C8")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], font_size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, "E8EEF5")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        run.font.size = Pt(font_size)
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for row_data in rows:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        prevent_row_split(row)
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            run.font.size = Pt(font_size)
    set_table_width(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Title": (30, DARK_BLUE, 0, 10),
        "Subtitle": (15, GRAY, 0, 12),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, "1F4D78", 10, 5),
    }
    for style_name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string(GRAY)

    capture_style = doc.styles.add_style("Capture Label", WD_STYLE_TYPE.PARAGRAPH)
    capture_style.font.name = "Calibri"
    capture_style.font.size = Pt(10)
    capture_style.font.bold = True
    capture_style.font.color.rgb = RGBColor.from_string(BLUE)
    capture_style.paragraph_format.space_before = Pt(8)
    capture_style.paragraph_format.space_after = Pt(3)

    small_style = doc.styles.add_style("Small Text", WD_STYLE_TYPE.PARAGRAPH)
    small_style.font.name = "Calibri"
    small_style.font.size = Pt(9)
    small_style.font.color.rgb = RGBColor.from_string(GRAY)
    small_style.paragraph_format.space_after = Pt(4)
    small_style.paragraph_format.line_spacing = 1.1

    for current_section in doc.sections:
        current_section.different_first_page_header_footer = True
        header = current_section.header
        paragraph = header.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run("PROMATEC LUGARTH  |  MANUAL DE USUARIO")
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(GRAY)
        footer = current_section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_paragraph.add_run("Uso interno  |  Julio 2026  |  Página ")
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor.from_string(GRAY)
        add_field(footer_paragraph.add_run(), "PAGE")

    doc.core_properties.title = "Manual de Usuario - Sistema de Nóminas PROMATEC LUGARTH"
    doc.core_properties.subject = "Operación integral del sistema de personal, asistencias, nóminas y seguridad"
    doc.core_properties.author = "PROMATEC LUGARTH"
    doc.core_properties.keywords = "manual, nómina, asistencias, empleados, seguridad, Firebase"


def add_cover(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(24)
    if PROMATEC_LOGO.exists():
        paragraph.add_run().add_picture(str(PROMATEC_LOGO), width=Inches(0.78))
    paragraph.add_run("     ")
    if LUGARTH_LOGO.exists():
        paragraph.add_run().add_picture(str(LUGARTH_LOGO), width=Inches(0.92))

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(38)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("MANUAL DE USUARIO")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Sistema de Nóminas, Asistencias y Personal")
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand.paragraph_format.space_before = Pt(12)
    brand.paragraph_format.space_after = Pt(32)
    run = brand.add_run("PROMATEC LUGARTH")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(TEAL)

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = line.add_run("Guía integral de operación, roles, seguridad y continuidad")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(GRAY)

    metadata = doc.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_before = Pt(80)
    metadata.paragraph_format.line_spacing = 1.4
    metadata.add_run("Versión documental: 1.0\n").bold = True
    metadata.add_run("Referencia funcional: julio de 2026\n")
    metadata.add_run("Documento editable con guía de 46 capturas\n")
    metadata.add_run("Clasificación: Uso interno")

    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    doc.add_heading("Control del documento", level=1)
    add_table(
        doc,
        ["Campo", "Detalle"],
        [
            ["Documento", "Manual de Usuario del Sistema de Nóminas, Asistencias y Personal"],
            ["Organización", "PROMATEC LUGARTH"],
            ["Versión", "1.0"],
            ["Fecha", "Julio de 2026"],
            ["Propietario sugerido", "Administración / Recursos Humanos"],
            ["Revisión sugerida", "Cada año y después de cambios funcionales importantes"],
        ],
        [2500, 6860],
        font_size=9,
    )

    doc.add_heading("Objetivo", level=2)
    doc.add_paragraph(
        "Este manual explica el uso completo del sistema web para administrar empleados, "
        "asistencias, incidencias, vacaciones, faltas, nóminas, diferencias IMSS, días festivos, "
        "respaldos, usuarios, permisos, auditoría y sincronización con la aplicación Mi Lugarth."
    )
    doc.add_heading("Audiencia", level=2)
    add_bullets(
        doc,
        [
            "Administradores responsables de configuración, seguridad, respaldos y continuidad.",
            "Gerencia y Recursos Humanos responsables de empleados, asistencias y nóminas.",
            "Capturistas responsables de incidencias e importación del reloj.",
            "Usuarios de consulta que revisan información sin modificarla.",
            "Personal de soporte que atiende accesos, sincronización y recuperación operativa.",
        ],
    )
    add_note(
        doc,
        "Alcance",
        "El manual describe la operación visible y las reglas ya implementadas. No autoriza modificar "
        "fórmulas, código, migraciones ni datos directamente en la base de datos.",
        "warn",
    )
    doc.add_heading("Cómo completar las capturas", level=2)
    doc.add_paragraph(
        "Los marcadores CAP-01 a CAP-42 corresponden al sistema web y APP-01 a APP-04 a la aplicación "
        "móvil. Sustituya cada recuadro por la captura indicada. Utilice siempre información ficticia "
        "o autorizada y oculte contraseñas, CLABE, RFC, CURP, NSS, teléfonos, correos e IP cuando no sean necesarios."
    )
    add_note(
        doc,
        "Formato recomendado",
        "Captura PNG, navegador al 100 %, resolución mínima 1366 × 768, menú visible cuando la guía lo pida "
        "y recorte horizontal sin barra de tareas. En móvil, use la captura nativa del dispositivo.",
        "good",
    )

    doc.add_heading("Contenido", level=1)
    sections = [
        "1. Acceso y ciclo de las cuentas",
        "2. Roles y permisos",
        "3. Navegación y panel principal",
        "4. Gestión de empleados",
        "5. Asistencias e importación del reloj",
        "6. Horas de alumnos",
        "7. Nóminas, recibos y exportaciones",
        "8. Días festivos",
        "9. Base de datos y respaldos",
        "10. Usuarios, permisos y auditoría",
        "11. Perfil, contraseña y cierre de sesión",
        "12. Aplicación Mi Lugarth y Firebase",
        "13. Reglas especiales y criterios operativos",
        "14. Rutinas recomendadas",
        "15. Solución de problemas",
        "16. Lista maestra de capturas",
        "17. Glosario y lista de entrega",
    ]
    add_bullets(doc, sections)
    doc.add_page_break()


def add_access_and_roles(doc: Document) -> None:
    doc.add_heading("1. Acceso y ciclo de las cuentas", level=1)
    doc.add_heading("1.1 Inicio de sesión", level=2)
    add_steps(
        doc,
        "Procedimiento",
        [
            "Abra la dirección del sistema proporcionada por el administrador.",
            "Capture el correo y la contraseña de su cuenta.",
            "Active Recordarme únicamente en un equipo personal o controlado.",
            "Seleccione Iniciar sesión.",
            "Si la cuenta está aprobada y activa, el sistema abrirá el Panel; de lo contrario mostrará el estado correspondiente.",
        ],
    )
    add_capture(doc, "CAP-01")
    add_note(
        doc,
        "Seguridad",
        "No comparta credenciales ni permita que el navegador guarde la contraseña en equipos compartidos.",
        "danger",
    )

    doc.add_heading("1.2 Registro y aprobación", level=2)
    doc.add_paragraph(
        "Una persona puede registrar una cuenta, pero el registro no concede acceso automático. "
        "Un administrador debe revisar la identidad, asignar el rol, aprobar la cuenta y guardar los cambios."
    )
    add_steps(
        doc,
        "Alta de cuenta",
        [
            "Desde la pantalla de acceso, abra Registrarse.",
            "Capture nombre, correo, contraseña y confirmación.",
            "Envíe el formulario.",
            "Espere la aprobación del administrador.",
            "El administrador entra a Seguridad > Usuarios, selecciona rol y permisos, marca Aprobado y guarda.",
        ],
    )
    add_capture(doc, "CAP-02")
    add_capture(doc, "CAP-03")

    doc.add_heading("1.3 Recuperación de contraseña", level=2)
    add_steps(
        doc,
        "Procedimiento",
        [
            "Seleccione ¿Olvidó su contraseña? en la pantalla de acceso.",
            "Capture el correo registrado y solicite el enlace.",
            "Abra el mensaje recibido y defina una contraseña nueva.",
            "Si no llega el correo, solicite al administrador que confirme que la cuenta y el correo sean correctos.",
        ],
    )
    add_note(
        doc,
        "Cuenta de recuperación",
        "El sistema protege una cuenta administrativa de recuperación: no puede deshabilitarse, eliminarse ni perder el rol Administrador.",
        "info",
    )

    doc.add_heading("2. Roles y permisos", level=1)
    doc.add_paragraph(
        "Los menús y acciones se muestran según los permisos efectivos. El rol aporta una configuración inicial; "
        "el administrador puede personalizar permisos para cada cuenta, excepto en Administrador, que conserva acceso total."
    )
    role_rows = [
        ["Ver panel", "Sí", "Sí", "Sí", "Sí"],
        ["Ver empleados", "Sí", "Sí", "Sí", "Sí"],
        ["Administrar empleados", "Sí", "Sí", "—", "—"],
        ["Ver asistencias", "Sí", "Sí", "Sí", "Sí"],
        ["Administrar asistencias", "Sí", "Sí", "Sí", "—"],
        ["Importar CSV del reloj", "Sí", "Sí", "Sí", "—"],
        ["Exportar asistencias", "Sí", "Sí", "Sí", "—"],
        ["Ver nóminas", "Sí", "Sí", "—", "Sí"],
        ["Administrar nóminas", "Sí", "Sí", "—", "—"],
        ["Marcar pago / revertir", "Sí", "Sí", "—", "—"],
        ["Exportar nóminas/PDF/Excel", "Sí", "Sí", "—", "—"],
        ["Días festivos", "Sí", "Sí", "—", "—"],
        ["Respaldos de base de datos", "Sí", "—", "—", "—"],
        ["Usuarios y permisos", "Sí", "—", "—", "—"],
        ["Auditoría", "Sí", "—", "—", "—"],
    ]
    add_table(
        doc,
        ["Permiso", "Administrador", "Gerente / RH", "Capturista", "Solo consulta"],
        role_rows,
        [4300, 1265, 1265, 1265, 1265],
        font_size=8,
    )
    doc.add_heading("2.1 Responsabilidad por rol", level=2)
    add_bullets(
        doc,
        [
            "Administrador: controla seguridad, respaldos, auditoría y todos los módulos operativos.",
            "Gerente / RH: opera empleados, asistencias, nóminas, pagos, exportaciones y días festivos.",
            "Capturista: captura, corrige, importa y exporta asistencias; consulta empleados.",
            "Solo consulta: revisa panel, empleados, asistencias y nóminas sin guardar cambios.",
        ],
    )
    add_note(
        doc,
        "Principio de mínimo acceso",
        "Asigne únicamente los permisos necesarios. Una persona que solo captura asistencias no debe recibir respaldos, usuarios ni pagos.",
        "good",
    )
    doc.add_page_break()


def add_navigation_dashboard(doc: Document) -> None:
    doc.add_heading("3. Navegación y panel principal", level=1)
    doc.add_heading("3.1 Estructura general", level=2)
    add_bullets(
        doc,
        [
            "Principal: Panel, Empleados, Asistencias y Nóminas.",
            "Sistema: Configuración, Días festivos y Base de datos.",
            "Seguridad: Usuarios y Auditoría.",
            "Barra superior: menú, estado del sistema, nombre, rol y Salir.",
        ],
    )
    doc.add_paragraph(
        "El menú se adapta al rol. Que un módulo no aparezca normalmente significa que la cuenta no tiene el permiso requerido."
    )
    add_capture(doc, "CAP-05")

    doc.add_heading("3.2 Panel principal", level=2)
    doc.add_paragraph(
        "El Panel concentra indicadores del periodo contable, operación de nómina, asistencia y personal. "
        "También muestra rankings de puntualidad e impuntualidad por semana, mes y año, distribución de incidencias, "
        "horas extra, cumpleaños y estado general."
    )
    add_bullets(
        doc,
        [
            "Use el Panel para detectar pendientes y anomalías; confirme siempre el detalle en el módulo correspondiente.",
            "Una semana sin registros debe mostrarse como ausencia de información, no como asistencia perfecta.",
            "Los rangos de semana, mes y año cambian automáticamente con la fecha seleccionada.",
            "Los nombres extensos se muestran completos o pueden consultarse en el detalle correspondiente.",
        ],
    )
    add_capture(doc, "CAP-04")
    add_note(
        doc,
        "Lectura correcta",
        "Los indicadores no sustituyen la revisión de la matriz semanal ni la previsualización CSV antes de generar nómina.",
        "warn",
    )


def add_employees(doc: Document) -> None:
    doc.add_heading("4. Gestión de empleados", level=1)
    doc.add_heading("4.1 Alta de empleado", level=2)
    doc.add_paragraph(
        "Entre a Empleados. El formulario de alta reúne información laboral, pago, deducciones, identificación y contacto."
    )
    add_steps(
        doc,
        "Procedimiento",
        [
            "Seleccione Registrar y asigne un número de empleado único.",
            "Capture nombre completo, puesto y fecha de ingreso.",
            "Seleccione Efectivo o Depósito.",
            "Indique si es alumno. Para alumno capture tarifa por hora; para personal regular capture sueldo base semanal.",
            "Registre deuda y descuento periódico de préstamo únicamente cuando correspondan.",
            "Abra Más datos y complete deducciones, banco, identificación y contactos.",
            "Revise la información y guarde.",
        ],
    )
    add_capture(doc, "CAP-06")
    add_capture(doc, "CAP-07")
    add_note(
        doc,
        "Validaciones",
        "El número de empleado no puede repetirse. Para pago por depósito se requiere banco y cuenta/CLABE. "
        "No use el número de una persona dada de baja sin revisar primero una posible restauración.",
        "warn",
    )

    doc.add_heading("4.2 Directorio, búsqueda y edición", level=2)
    add_steps(
        doc,
        "Consultar o modificar",
        [
            "Use la búsqueda por nombre o número.",
            "Seleccione el filtro Activos, Bajas/Papelera o Préstamo según la revisión.",
            "Abra Perfil para consultar el expediente.",
            "Use Editar para modificar datos y guarde una vez confirmados.",
        ],
    )
    add_capture(doc, "CAP-08")

    doc.add_heading("4.3 Fotografías", level=2)
    doc.add_paragraph(
        "Seleccione la fotografía para abrir el visor. Puede cerrarlo con el botón de cierre, haciendo clic fuera de la imagen "
        "o con la tecla Escape. La fotografía debe corresponder al número y expediente del empleado."
    )
    add_capture(doc, "CAP-09")
    add_note(
        doc,
        "Control de imagen",
        "Al dar de baja, el sistema clasifica la fotografía en bajas; al restaurar, la devuelve al grupo activo. "
        "Si la foto falta o no corresponde, no sustituya archivos manualmente: repórtelo a soporte con el número e ID del empleado.",
        "info",
    )

    doc.add_heading("4.4 Perfil integral", level=2)
    doc.add_paragraph(
        "El Perfil centraliza datos personales, condiciones de nómina, préstamo, asistencia, vacaciones, faltas y acceso móvil."
    )
    add_capture(doc, "CAP-10")
    add_capture(doc, "CAP-11")
    add_capture(doc, "CAP-12")

    doc.add_heading("4.5 Baja de empleado", level=2)
    add_steps(
        doc,
        "Procedimiento",
        [
            "Localice al empleado activo y seleccione Dar de baja.",
            "Capture la fecha efectiva y el motivo.",
            "Verifique que la fecha sea correcta antes de confirmar.",
            "Confirme la baja y abra el perfil para validar el resultado.",
        ],
    )
    doc.add_paragraph(
        "La baja conserva el historial. El perfil muestra los días laborados desde el ingreso hasta la baja y, adicionalmente, "
        "los días laborados dentro del año de la baja. El conteo excluye domingos y se adapta al año registrado."
    )
    add_capture(doc, "CAP-13")
    add_capture(doc, "CAP-14")

    doc.add_heading("4.6 Restauración", level=2)
    add_steps(
        doc,
        "Procedimiento",
        [
            "Abra el filtro Bajas/Papelera.",
            "Localice al empleado y seleccione Restaurar.",
            "Si su número anterior está libre, el sistema lo recupera.",
            "Si el número ya pertenece a otro empleado activo, la persona se restaura sin número.",
            "Edite de inmediato el expediente y asigne un número único antes de usar reloj, asistencia o nómina.",
        ],
    )
    add_capture(doc, "CAP-15")
    add_note(
        doc,
        "Mensaje esperado",
        "Cuando el número anterior esté ocupado, el sistema avisa que debe asignarse uno nuevo antes de utilizar el checador.",
        "warn",
    )

    doc.add_heading("4.7 Acceso a Mi Lugarth", level=2)
    add_steps(
        doc,
        "Activar acceso",
        [
            "Abra el perfil del empleado.",
            "En Acceso móvil, capture un nombre de usuario y una contraseña temporal.",
            "Guarde el acceso y entregue las credenciales por un medio seguro.",
            "Pida al empleado validar su información en la aplicación.",
            "Use Desactivar acceso cuando la cuenta ya no deba entrar.",
        ],
    )
    add_capture(doc, "CAP-16")
    doc.add_page_break()


def add_attendance(doc: Document) -> None:
    doc.add_heading("5. Asistencias e importación del reloj", level=1)
    doc.add_heading("5.1 Estructura del módulo", level=2)
    add_bullets(
        doc,
        [
            "Captura y Reloj: carga CSV y registro manual.",
            "Revisión CSV: valida, corrige y aprueba la previsualización.",
            "Control Vacaciones: consulta saldos y consumos.",
            "Control Faltas: consulta faltas por año y empleado.",
            "Horas de alumnos: abre el control semanal y su PDF.",
        ],
    )

    doc.add_heading("5.2 Captura manual", level=2)
    add_steps(
        doc,
        "Registrar asistencia o incidencia",
        [
            "Busque y seleccione al empleado.",
            "Indique la fecha.",
            "Seleccione Normal, Falta, Incapacidad o Vacaciones.",
            "Para Normal, capture entrada y salida; para una incidencia confirme el tipo correcto.",
            "Revise horas, retraso y demás indicadores calculados.",
            "Guarde. Para corregir, use Editar en la matriz semanal.",
        ],
    )
    doc.add_paragraph(
        "Las incidencias disponibles incluyen Falta injustificada, Incapacidad al 60 % y Vacaciones con prima del 25 %."
    )
    add_capture(doc, "CAP-18")

    doc.add_heading("5.3 Control semanal", level=2)
    doc.add_paragraph(
        "La semana operativa se presenta de jueves a miércoles. Use las flechas o la fecha de referencia para cambiar de periodo; "
        "busque por nombre/número y revise por día el tipo, entrada, salida, horas normales, extra y retraso."
    )
    add_capture(doc, "CAP-19")

    doc.add_heading("5.4 Cargar archivo CSV del reloj", level=2)
    add_steps(
        doc,
        "Previsualizar",
        [
            "Abra Captura y Reloj.",
            "Seleccione el archivo CSV o TXT del reloj, con tamaño máximo de 10 MB.",
            "Si solo necesita un periodo específico, capture Inicio semana y Fin semana.",
            "Envíe el archivo para generar la previsualización.",
            "No cierre el proceso hasta revisar el resumen y corregir las filas necesarias.",
        ],
    )
    add_capture(doc, "CAP-17")
    add_note(
        doc,
        "Rango de fechas",
        "Cuando se captura un rango en el sistema, ese rango tiene prioridad y se ignoran marcaciones fuera de él. "
        "Si se dejan las fechas vacías, se usa el mínimo y máximo detectados en el CSV.",
        "info",
    )
    add_note(
        doc,
        "Marcaciones duplicadas",
        "Las marcaciones repetidas dentro de aproximadamente dos minutos se consolidan para evitar entradas o salidas duplicadas.",
        "good",
    )

    doc.add_heading("5.5 Revisión CSV", level=2)
    doc.add_paragraph(
        "La revisión no guarda asistencias hasta seleccionar y aprobar. El resumen identifica registros nuevos, actualizaciones, "
        "faltas generadas, incompletas y números que no corresponden a un empleado."
    )
    add_capture(doc, "CAP-20")
    add_bullets(
        doc,
        [
            "Sin registro: no hubo marcaciones en un día laboral esperado; puede aprobarse como Falta o editarse.",
            "Incompleta: falta entrada o salida; complete la hora antes de aprobar.",
            "Actualiza: ya existía asistencia para ese día y la aprobación reemplazará los datos.",
            "Existente: coincide con información ya registrada.",
            "Sin empleado: el número del CSV no pudo relacionarse; corrija el empleado antes de aprobar.",
        ],
    )

    doc.add_heading("5.6 Corregir incompletas y faltas", level=2)
    add_steps(
        doc,
        "Corrección segura",
        [
            "Seleccione el filtro Incompletas para trabajar únicamente los registros que requieren una hora.",
            "Corrija empleado, fecha, tipo, entrada o salida según la evidencia.",
            "Cambie de página cuando sea necesario; las ediciones se conservan en la sesión de revisión.",
            "Use el filtro Faltas para revisar días sin marcación y convertirlos a otra incidencia cuando exista soporte.",
            "Desmarque cualquier fila que no deba guardarse.",
        ],
    )
    add_capture(doc, "CAP-21")
    add_capture(doc, "CAP-22")
    add_note(
        doc,
        "No inventar horarios",
        "Si una persona olvidó entrada o salida, valide la hora con el responsable antes de completar el registro.",
        "danger",
    )

    doc.add_heading("5.7 Aprobar o descartar", level=2)
    add_steps(
        doc,
        "Cierre de revisión",
        [
            "Confirme que el rango sea el solicitado.",
            "Revise especialmente Incompletas, Faltas, Actualiza y Sin empleado.",
            "Seleccione únicamente las filas correctas.",
            "Pulse Aprobar seleccionadas y espere el mensaje de finalización.",
            "Use Descartar revisión si el archivo o el rango fueron incorrectos; después cargue nuevamente.",
        ],
    )
    add_capture(doc, "CAP-23")
    add_note(
        doc,
        "Rendimiento",
        "La aprobación procesa por bloques para evitar tiempos de espera extensos. No pulse Aprobar repetidamente; espere la respuesta.",
        "good",
    )

    doc.add_heading("5.8 Vacaciones y faltas", level=2)
    doc.add_paragraph(
        "Control Vacaciones muestra días totales, utilizados y restantes. Control Faltas permite seleccionar cualquier año "
        "disponible y desplegar las fechas por empleado."
    )
    add_capture(doc, "CAP-24")
    add_capture(doc, "CAP-25")
    doc.add_page_break()


def add_students(doc: Document) -> None:
    doc.add_heading("6. Horas de alumnos", level=1)
    doc.add_paragraph(
        "Este módulo muestra únicamente empleados marcados como alumnos y permite revisar sus horas por semana."
    )
    add_steps(
        doc,
        "Generar reporte",
        [
            "Abra Asistencias > Horas de alumnos.",
            "Seleccione la semana y use la búsqueda cuando sea necesario.",
            "Revise el número de empleado, las fechas y las horas.",
            "Seleccione los alumnos que formarán parte del reporte.",
            "Genere el PDF y confirme que cada hoja contenga dos registros.",
        ],
    )
    add_capture(doc, "CAP-26")
    add_capture(doc, "CAP-27")
    add_note(
        doc,
        "Alumnos",
        "El tiempo adicional de los alumnos se considera dentro de sus horas de servicio; no se trata como hora extra de personal regular.",
        "info",
    )


def add_payroll(doc: Document) -> None:
    doc.add_heading("7. Nóminas, recibos y exportaciones", level=1)
    doc.add_heading("7.1 Antes de procesar", level=2)
    add_bullets(
        doc,
        [
            "Confirme que la semana de jueves a miércoles sea la correcta.",
            "Apruebe o capture todas las asistencias laborales esperadas.",
            "Revise faltas, vacaciones, incapacidades, retardos y horas extra.",
            "Verifique sueldo/tarifa, deducciones, banco, préstamo e IMSS en el expediente.",
            "Genere primero recibos de prueba cuando existan cambios excepcionales.",
        ],
    )
    add_note(
        doc,
        "Asistencia pendiente",
        "El sistema puede bloquear o advertir el cálculo cuando faltan asistencias esperadas. Corrija el módulo Asistencias antes de continuar.",
        "warn",
    )

    doc.add_heading("7.2 Periodo, filtros y selección", level=2)
    doc.add_paragraph(
        "Seleccione la semana y filtre por Todos, Pendientes o Liquidados. Puede ordenar, buscar y filtrar por banco. "
        "Las acciones masivas respetan la selección actual."
    )
    add_capture(doc, "CAP-28")
    add_capture(doc, "CAP-29")

    doc.add_heading("7.3 Tarjeta de empleado", level=2)
    doc.add_paragraph(
        "Cada tarjeta reúne estado, cuenta bancaria, pago neto, deuda, horas pendientes, ajustes y acciones de generación."
    )
    add_capture(doc, "CAP-30")

    doc.add_heading("7.4 Préstamo y ajustes", level=2)
    add_bullets(
        doc,
        [
            "Entregar hoy aumenta el préstamo otorgado en el periodo.",
            "Descontar hoy aplica la parcialidad autorizada.",
            "Descuento manual registra un ajuste extraordinario documentado.",
            "Días adicionales de vacaciones ajustan el periodo cuando existe autorización.",
            "Los ajustes avanzados permiten revisar faltas pagadas/adeudadas, vacaciones, incapacidad y horas extra detectadas/por pagar.",
        ],
    )
    add_capture(doc, "CAP-31")
    add_capture(doc, "CAP-33")
    add_note(
        doc,
        "Trazabilidad",
        "Antes de guardar un ajuste manual, conserve el soporte autorizado. Los cambios quedan asociados a la nómina y pueden aparecer en auditoría.",
        "warn",
    )

    doc.add_heading("7.5 Diferencia IMSS", level=2)
    doc.add_paragraph(
        "Depósito IMSS, Diferencia semanal y Suma total pertenecen a un control independiente. "
        "La diferencia no modifica el formato del recibo individual normal."
    )
    add_capture(doc, "CAP-32")

    doc.add_heading("7.6 Generar recibos", level=2)
    add_steps(
        doc,
        "Recibo normal",
        [
            "Revise los datos y ajustes del empleado.",
            "Seleccione Generar o Regenerar.",
            "Abra el PDF y confirme nombre, número, periodo, percepciones, deducciones y neto.",
            "Use PDF seleccionados, PDF todos o PDF grupo para impresión masiva.",
            "Confirme que la salida masiva conserve dos recibos por hoja.",
        ],
    )
    add_capture(doc, "CAP-34")
    add_capture(doc, "CAP-35")

    doc.add_heading("7.7 Recibos de diferencias IMSS", level=2)
    doc.add_paragraph(
        "PDF Diferencias IMSS genera recibos aparte únicamente para empleados con una diferencia registrada distinta de cero. "
        "El recibo usa el concepto Sueldo para mostrar la diferencia; los demás importes permanecen en cero y la salida masiva "
        "conserva dos recibos por hoja."
    )
    add_steps(
        doc,
        "Generar diferencias",
        [
            "Confirme y guarde la diferencia IMSS en cada empleado aplicable.",
            "Seleccione PDF Diferencias IMSS en la barra superior.",
            "Revise que solo aparezcan empleados con diferencia.",
            "Valide el importe y el formato antes de imprimir.",
        ],
    )
    add_capture(doc, "CAP-36")

    doc.add_heading("7.8 Excel e historial", level=2)
    add_bullets(
        doc,
        [
            "Excel global exporta el resumen del periodo.",
            "Excel IMSS exporta la información correspondiente al control IMSS.",
            "El botón Excel de la tarjeta genera el archivo individual.",
            "Historial de recibos permite buscar semanas anteriores, revisar estado y volver a abrir el PDF.",
        ],
    )
    add_capture(doc, "CAP-37")

    doc.add_heading("7.9 Marcar pagada o revertir", level=2)
    add_steps(
        doc,
        "Liquidar",
        [
            "Revise el PDF y obtenga la autorización de pago.",
            "Seleccione Marcar pagada o use la acción masiva sobre la selección correcta.",
            "Confirme el estado Liquidada.",
            "Verifique la sincronización con Mi Lugarth cuando corresponda.",
        ],
    )
    doc.add_paragraph(
        "Al liquidar se aplican movimientos asociados, como préstamo y vacaciones. Revertir devuelve la nómina a pendiente "
        "y revierte esos movimientos relacionados. Use la reversión solo para corregir un pago marcado por error."
    )
    add_note(
        doc,
        "Cierre",
        "No cambie asistencias ni deducciones después de pagar sin revisar primero el impacto y, cuando proceda, revertir la nómina.",
        "danger",
    )
    doc.add_page_break()


def add_system_admin(doc: Document) -> None:
    doc.add_heading("8. Días festivos", level=1)
    doc.add_paragraph(
        "Los días festivos activos influyen en asistencia, generación de faltas, panel y nómina. "
        "El módulo admite años de 2024 a 2100 y ofrece años cercanos para operación rápida."
    )
    add_steps(
        doc,
        "Preparar un año",
        [
            "Seleccione el año.",
            "Pulse Generar oficiales de México.",
            "Revise fechas y active o desactive según la operación de la empresa.",
            "Agregue manualmente días empresariales, electorales u otros cuando correspondan.",
            "No duplique una fecha existente.",
        ],
    )
    add_capture(doc, "CAP-38")
    add_note(
        doc,
        "Inicio de año",
        "Genere y valide los días festivos antes de importar asistencias o procesar la primera nómina del nuevo año.",
        "good",
    )

    doc.add_heading("9. Base de datos y respaldos", level=1)
    doc.add_paragraph(
        "Este módulo es exclusivo de Administrador y funciona con la base MySQL configurada. "
        "Muestra conexión, nombre de la base y tablas detectadas."
    )
    add_steps(
        doc,
        "Exportar respaldo",
        [
            "Abra Sistema > Base de datos.",
            "Confirme que la conexión y el nombre sean los esperados.",
            "Seleccione Exportar respaldo.",
            "Guarde el archivo SQL en una ubicación externa y protegida.",
            "Registre fecha, responsable y motivo del respaldo.",
        ],
    )
    add_steps(
        doc,
        "Restaurar",
        [
            "Exporte primero un respaldo del estado actual.",
            "Seleccione un archivo SQL o TXT generado por el sistema, con tamaño máximo de 100 MB.",
            "Escriba exactamente RESTAURAR en la confirmación.",
            "Inicie el proceso y no cierre el navegador.",
            "Al finalizar, valide acceso, empleados, asistencias, nóminas y auditoría.",
        ],
    )
    add_capture(doc, "CAP-39")
    add_note(
        doc,
        "Riesgo crítico",
        "La restauración reemplaza la información actual. No restaure una base recibida de otra computadora sin respaldar y confirmar versión, entorno y archivo.",
        "danger",
    )
    add_note(
        doc,
        "Tabla existente",
        "Si una migración informa que audit_logs ya existe, no continúe borrando tablas manualmente. "
        "Use un respaldo compatible y solicite a soporte revisar el historial de migraciones de esa instalación.",
        "warn",
    )

    doc.add_heading("10. Usuarios, permisos y auditoría", level=1)
    doc.add_heading("10.1 Usuarios y permisos", level=2)
    add_steps(
        doc,
        "Aprobar o modificar una cuenta",
        [
            "Abra Seguridad > Usuarios.",
            "Revise nombre, correo, estado, último acceso e IP.",
            "Seleccione el rol.",
            "Marque Aprobado o Deshabilitado según corresponda.",
            "Abra Permisos asignados y ajuste únicamente lo necesario.",
            "Guarde y confirme el estado final.",
        ],
    )
    add_capture(doc, "CAP-40")
    add_bullets(
        doc,
        [
            "No puede quitarse a sí mismo el acceso administrativo.",
            "Siempre debe quedar al menos un administrador aprobado y activo.",
            "La cuenta de recuperación no puede deshabilitarse ni dejar de ser Administrador.",
            "Eliminar una cuenta es diferente de deshabilitarla; para una baja temporal prefiera Deshabilitado.",
        ],
    )

    doc.add_heading("10.2 Auditoría", level=2)
    doc.add_paragraph(
        "Auditoría registra acciones sensibles, cambios, exportaciones y descargas. Puede filtrarse por texto, usuario, acción y rango de fechas."
    )
    add_steps(
        doc,
        "Investigar una acción",
        [
            "Defina usuario, acción o fechas.",
            "Abra el registro relevante.",
            "Revise descripción, valores anteriores/nuevos y detalle técnico.",
            "Conserve evidencia antes de eliminar registros.",
            "Use Purgar únicamente conforme a la política de retención.",
        ],
    )
    add_capture(doc, "CAP-41")
    add_note(
        doc,
        "Privacidad",
        "La auditoría oculta contraseñas y tokens. Los correos pueden mostrarse limitados a usuarios no administradores.",
        "info",
    )

    doc.add_heading("11. Perfil, contraseña y cierre de sesión", level=1)
    add_bullets(
        doc,
        [
            "Configuración permite actualizar nombre y correo.",
            "Cambio de contraseña requiere la contraseña actual y confirmación de la nueva.",
            "Eliminar cuenta es definitivo y puede estar protegido para cuentas especiales.",
            "Salir finaliza la sesión; úselo siempre al terminar en un equipo compartido.",
        ],
    )
    add_capture(doc, "CAP-42")
    doc.add_page_break()


def add_mobile_and_rules(doc: Document) -> None:
    doc.add_heading("12. Aplicación Mi Lugarth y Firebase", level=1)
    doc.add_paragraph(
        "El acceso móvil se vincula desde el perfil del empleado. La integración sincroniza información del empleado, "
        "asistencias y nóminas pagadas para consulta en Mi Lugarth."
    )
    add_steps(
        doc,
        "Validación de una cuenta móvil",
        [
            "Active el acceso desde el perfil y entregue credenciales temporales.",
            "Inicie sesión en Mi Lugarth con el empleado de prueba.",
            "Confirme nombre, número y resumen.",
            "Capture o importe una asistencia en web y verifique que aparezca en la aplicación.",
            "Marque una nómina como pagada y verifique que se incorpore al historial móvil.",
        ],
    )
    add_capture(doc, "APP-01")
    add_capture(doc, "APP-02")
    add_capture(doc, "APP-03")
    add_capture(doc, "APP-04")
    add_note(
        doc,
        "Sincronización",
        "Las asistencias nuevas, editadas o importadas y las nóminas pagadas se envían a Firebase. "
        "Si no se reflejan, no duplique registros: valide conexión y solicite una resincronización.",
        "warn",
    )

    doc.add_heading("13. Reglas especiales y criterios operativos", level=1)
    rules = [
        ["Semana de nómina", "Jueves a miércoles."],
        ["Día laboral regular", "Lunes a viernes. Sábado y domingo no generan falta para personal regular."],
        ["Trabajo en fin de semana", "Se considera tiempo extra para personal regular cuando existen marcaciones válidas."],
        ["Hora extra entre semana", "Se reconoce después de las 17:30 en bloques completos de media hora."],
        ["Hora extra fin de semana", "Se redondea al bloque de media hora más cercano, a partir de las 08:00."],
        ["Media hora", "Las horas extra pueden expresarse como 0.5; no se limitan a horas completas."],
        ["Alumnos", "No generan retardo ni hora extra; el tiempo se acumula como horas de servicio."],
        ["Vigilancia / Seguridad", "Jornada 24 × 24, sin retardos ni horas extra; faltas según su rotación."],
        ["Detección vigilancia", "Primero por puesto que contenga VIGILANCIA o SEGURIDAD; respaldo para empleados 20 y 29."],
        ["Sin horas extra", "Empleados 8, 9 y 22, además de Vigilancia/Seguridad."],
        ["Sin retardo", "Empleados 14, 76 y 78, además de Vigilancia/Seguridad."],
        ["Tope de horas", "Empleados 76 y 78: hasta 48 horas para el criterio configurado."],
        ["Falta regular", "Solo se genera en día laboral esperado, excluyendo fines de semana y festivos activos."],
        ["Marcación incompleta", "Se conserva para corrección; no debe convertirse automáticamente en horario inventado."],
        ["Días laborados en baja", "Desde ingreso hasta baja, excluyendo domingos; también se informa el total del año de baja."],
        ["Días festivos", "Deben generarse y validarse cada año antes de operar ese periodo."],
    ]
    add_table(doc, ["Regla", "Comportamiento vigente"], rules, [2800, 6560], font_size=8)
    add_note(
        doc,
        "Puestos especiales",
        "Para futuros vigilantes, registre correctamente el puesto. Usar el puesto es más escalable que depender solamente del número de empleado.",
        "good",
    )

    doc.add_heading("13.1 Continuidad hacia años futuros", level=2)
    add_bullets(
        doc,
        [
            "Los periodos y conteos se basan en fechas completas, por lo que cambian de año sin reiniciar datos.",
            "El control de faltas permite seleccionar el año y el perfil de baja calcula también el año de la baja.",
            "Los días festivos pueden administrarse hasta 2100.",
            "Al comenzar cada año, genere festivos, revise calendarios, valide la primera semana y cree un respaldo.",
            "Mantenga puestos, fechas de ingreso y estatus correctamente capturados para conservar cálculos y reportes.",
        ],
    )


def add_operations_troubleshooting(doc: Document) -> None:
    doc.add_heading("14. Rutinas recomendadas", level=1)
    doc.add_heading("14.1 Rutina diaria", level=2)
    add_bullets(
        doc,
        [
            "Revisar marcaciones incompletas y confirmar olvidos de entrada o salida.",
            "Capturar incidencias autorizadas.",
            "Validar altas, bajas, restauraciones y fotografías.",
            "Revisar alertas de sincronización móvil.",
        ],
    )
    doc.add_heading("14.2 Cierre semanal", level=2)
    add_steps(
        doc,
        "Orden recomendado",
        [
            "Respaldar la base antes de operaciones masivas o excepcionales.",
            "Importar el CSV usando el rango solicitado.",
            "Corregir Incompletas, Faltas, Actualiza y Sin empleado.",
            "Aprobar y revisar la matriz semanal.",
            "Validar vacaciones, incapacidades, retardos y horas extra.",
            "Procesar ajustes, préstamos e IMSS.",
            "Generar y revisar recibos/Excel.",
            "Marcar pagadas únicamente las nóminas autorizadas.",
            "Confirmar la sincronización móvil.",
        ],
    )
    doc.add_heading("14.3 Inicio de mes y año", level=2)
    add_bullets(
        doc,
        [
            "Exportar un respaldo y conservarlo fuera del equipo principal.",
            "Revisar cuentas activas, pendientes y deshabilitadas.",
            "Revisar auditoría y política de retención.",
            "Al iniciar año, generar festivos, validar filtros y probar una semana completa antes del primer pago.",
        ],
    )

    doc.add_heading("15. Solución de problemas", level=1)
    troubleshooting = [
        ["Pantalla en blanco al abrir por IP", "Ejecute la compilación de producción, confirme que los archivos de public/build se copiaron, revise APP_URL/Vite, reinicie el servidor y permita el puerto en el firewall. No dependa del servidor de desarrollo en la computadora principal."],
        ["Un módulo no aparece", "Revise el rol y permisos efectivos en Seguridad > Usuarios; cierre sesión y vuelva a entrar después del cambio."],
        ["Cuenta pendiente o deshabilitada", "Un Administrador debe aprobar/reactivar la cuenta y guardar."],
        ["CSV detecta en blanco", "Confirme rango, número de empleado, formato de fecha y clasificación Sin empleado/Sin registro. Revise la semana correcta."],
        ["Entrada o salida faltante", "Use filtro Incompletas, valide la hora con el responsable y complete solo con evidencia."],
        ["Edición desaparece al cambiar de página", "Permanezca en la misma revisión y navegador. No descarte ni cargue otro CSV; vuelva a la página y confirme antes de aprobar."],
        ["Aprobación excede el tiempo", "Espere la respuesta sin repetir. Si falla, conserve la revisión, revise el registro del servidor y vuelva a intentar por una selección menor."],
        ["Domingo no cuenta como extra", "Confirme entrada y salida válidas, que no sea alumno/Vigilancia ni un empleado excluido, y regenere/revise la nómina después de guardar asistencia."],
        ["Media hora extra no aparece", "Confirme que el tramo cumpla la regla de 30 minutos y que el empleado no esté excluido. Revise 0.5 en asistencia y luego en nómina/app."],
        ["Nómina usa sueldo semanal completo", "Verifique que toda la semana tenga asistencias o faltas aprobadas; regenere el recibo después de corregir la asistencia."],
        ["Foto incorrecta o desaparecida", "No renombre ni mueva archivos manualmente. Reporte número e ID; valide baja/restauración y la foto asociada."],
        ["Número duplicado al restaurar", "El sistema restaura sin número cuando el anterior está ocupado. Edite y asigne uno único antes del checador."],
        ["Mi Lugarth no actualiza", "Revise conexión, acceso móvil activo y que la nómina esté pagada. Solicite sincronización completa sin duplicar datos."],
        ["Error audit_logs ya existe", "No borre repetidamente la tabla. Respalde y haga que soporte alinee la tabla con el historial de migraciones."],
        ["Restauración de base falla", "Verifique MySQL, archivo del sistema, tamaño menor a 100 MB y confirmación RESTAURAR. No mezcle bases de entornos distintos sin revisión."],
    ]
    add_table(doc, ["Situación", "Qué revisar / acción"], troubleshooting, [2900, 6460], font_size=8)
    add_note(
        doc,
        "Escalamiento",
        "Al pedir soporte incluya fecha/hora, usuario, módulo, empleado, semana, acción realizada, mensaje exacto y captura sin datos sensibles.",
        "info",
    )
    doc.add_page_break()


def add_capture_checklist(doc: Document) -> None:
    doc.add_heading("16. Lista maestra de capturas", level=1)
    doc.add_paragraph(
        "Tome las capturas en este orden para reutilizar una sola cuenta de prueba y reducir cambios de contexto. "
        "Después sustituya cada marcador del manual por la imagen correspondiente."
    )
    rows = [[item["id"], item["screen"], item["state"], item["focus"]] for item in CAPTURES]
    chunk_size = 8
    chunks = [rows[index : index + chunk_size] for index in range(0, len(rows), chunk_size)]
    for chunk_index, chunk in enumerate(chunks):
        if chunk_index:
            doc.add_heading("16. Lista maestra de capturas (continuación)", level=2)
        table = add_table(
            doc,
            ["ID", "Pantalla", "Preparación", "Elementos obligatorios"],
            chunk,
            [900, 2100, 2800, 3560],
            font_size=7,
        )
        for row in table.rows[1:]:
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(BLUE)
        doc.add_page_break()

    doc.add_heading("16.1 Datos de prueba sugeridos", level=2)
    add_bullets(
        doc,
        [
            "Una cuenta Administrador y una cuenta pendiente.",
            "Un empleado activo regular con depósito y préstamo.",
            "Un empleado dado de baja cuyo número anterior esté ocupado por otro registro de prueba.",
            "Un alumno con horas registradas.",
            "Un puesto de Vigilancia/Seguridad con rotación 24 × 24.",
            "Una semana con registro normal, falta, vacaciones, incapacidad, incompleta, domingo trabajado y 0.5 hora extra.",
            "Una nómina pendiente, una liquidada y una diferencia IMSS distinta de cero.",
        ],
    )
    add_note(
        doc,
        "Privacidad",
        "No use datos reales para preparar el manual final. Cuando sea inevitable, desenfoque o cubra identificadores y datos bancarios antes de insertar la imagen.",
        "danger",
    )
    doc.add_page_break()


def add_glossary(doc: Document) -> None:
    doc.add_heading("17. Glosario y lista de entrega", level=1)
    glossary = [
        ["Asistencia incompleta", "Registro con entrada o salida faltante que requiere validación y corrección."],
        ["Diferencia IMSS", "Importe separado del recibo normal que genera su propio comprobante."],
        ["Firebase", "Servicio que recibe la información sincronizada para Mi Lugarth."],
        ["Falta", "Ausencia registrada o generada en un día laboral esperado."],
        ["Horas extra", "Tiempo adicional reconocido por las reglas del empleado y expresado también en medias horas."],
        ["Liquidada", "Nómina marcada como pagada, con movimientos relacionados aplicados."],
        ["Periodo contable", "Semana de nómina de jueves a miércoles."],
        ["Preview / Revisión CSV", "Área temporal donde se corrigen datos antes de guardarlos."],
        ["Rol", "Conjunto inicial de permisos asignado a una cuenta."],
        ["Sin empleado", "Número del CSV que no pudo relacionarse con un expediente activo."],
        ["Vigilancia 24 × 24", "Jornada especial alternada, sin retardo ni hora extra y con faltas según rotación."],
    ]
    glossary_chunks = [glossary[:6], glossary[6:]]
    for chunk_index, chunk in enumerate(glossary_chunks):
        if chunk_index:
            doc.add_heading("17. Glosario (continuación)", level=2)
        add_table(doc, ["Término", "Definición"], chunk, [2600, 6760], font_size=9)
        doc.add_page_break()

    doc.add_heading("17.1 Lista de entrega del manual", level=2)
    checklist = [
        "Se sustituyeron los 42 marcadores web y los 4 marcadores móviles.",
        "No aparecen contraseñas ni datos personales/bancarios sin protección.",
        "Las capturas corresponden a la versión instalada en la computadora de la contadora.",
        "Se revisaron nombres de botones, periodos y roles.",
        "Se validó una impresión de recibos normales, diferencias IMSS y horas de alumnos.",
        "Se confirmó la operación con un usuario distinto de Administrador.",
        "El archivo final se guardó en Word y PDF.",
        "Se asignó responsable y fecha de próxima revisión.",
    ]
    for item in checklist:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.2)
        paragraph.paragraph_format.space_after = Pt(5)
        run = paragraph.add_run("☐ ")
        run.font.name = "Segoe UI Symbol"
        run.font.size = Pt(12)
        paragraph.add_run(item)

    doc.add_heading("17.2 Aprobación", level=2)
    add_table(
        doc,
        ["Responsable", "Nombre", "Firma", "Fecha"],
        [
            ["Elaboró", "", "", ""],
            ["Revisó", "", "", ""],
            ["Autorizó", "", "", ""],
        ],
        [1800, 3000, 2760, 1800],
        font_size=9,
    )

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(28)
    run = paragraph.add_run("FIN DEL MANUAL")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string(TEAL)


def build_document() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_front_matter(doc)
    add_access_and_roles(doc)
    add_navigation_dashboard(doc)
    add_employees(doc)
    add_attendance(doc)
    add_students(doc)
    add_payroll(doc)
    add_system_admin(doc)
    add_mobile_and_rules(doc)
    add_operations_troubleshooting(doc)
    add_capture_checklist(doc)
    add_glossary(doc)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
